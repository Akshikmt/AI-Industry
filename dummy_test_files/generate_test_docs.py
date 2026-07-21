import os
import sys
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import docx
from PIL import Image, ImageDraw, ImageFont

OUTPUT_DIR = r"e:\Akshita Kumawat\ai industry\dummy_test_files"
os.makedirs(OUTPUT_DIR, exist_ok=True)

print(f"Generating multi-page industrial test documents in {OUTPUT_DIR}...")

# ---------------------------------------------------------
# 1. Multi-page TXT Document
# ---------------------------------------------------------
txt_path = os.path.join(OUTPUT_DIR, "SOP_505_Steam_Turbine_Overhaul_Manual.txt")
txt_content = """================================================================================
STANDARD OPERATING PROCEDURE: STEAM TURBINE T-202 OVERHAUL & CALIBRATION
Document ID: SOP-505-REV3 | Classification: Restricted Operational Document
Target Equipment: High-Pressure Steam Turbine (Tag: T-202) | Department: Maintenance
================================================================================

--- PAGE 1 OF 2 ---

1. OPERATIONAL OVERVIEW & SAFETY MANDATES
--------------------------------------------------------------------------------
This document establishes standard operating protocols for routine calibration, 
mechanical governor inspection, and bearing vibration alignment for High Pressure
Steam Turbine T-202.

SAFETY REQUIREMENTS:
- Lockout / Tagout (LOTO) protocol LOTO-881 must be strictly executed before access.
- All high-pressure steam lines must be isolated and depressurized to 0 PSIG.
- Personal Protective Equipment (PPE): Thermal protective suit, safety boots, face shield.

EQUIPMENT DESIGN SPECIFICATIONS (TAG: T-202):
- Max Operating Speed: 3,600 RPM
- Critical Temperature Limit: 85.0°C (Warning Trigger at 78.0°C)
- Vibration Threshold Limit: 2.5 mm/s RMS (Shutdown Alarm at 3.8 mm/s)
- Lube Oil Specification: ISO VG 46 Turbine Oil (Capacity: 120 Liters)

2. PRE-MAINTENANCE CHECKLIST & INSTRUMENT DIAGNOSTICS
--------------------------------------------------------------------------------
Step 1.1: Verify oil sump level indicator on T-202 sight glass.
Step 1.2: Check differential pressure across secondary lube oil filters (Max Delta P = 12 PSI).
Step 1.3: Calibrate proximity probes PT-202A and PT-202B on main rotor bearing shaft.
Step 1.4: Confirm emergency trip throttle valve solenoid responds within 250 milliseconds.


--- PAGE 2 OF 2 ---

3. ROTOR DECELERATION & GOVERNOR VALVE CALIBRATION PROCEDURE
--------------------------------------------------------------------------------
During deceleration phase tests, vibration harmonic spikes must be recorded.

DETAILED STEP-BY-STEP CALIBRATION:
Step 2.1: Disengage main generator breaker to isolate turbine T-202 from electrical grid.
Step 2.2: Gradually close main steam inlet throttle valve at a rate of 5% per minute.
Step 2.3: Monitor real-time vibration telemetry on AI Copilot dashboard.
         If vibration exceeds 3.2 mm/s RMS, abort test and execute manual trip switch.
Step 2.4: Inspect governor linkage pins for mechanical wear or corrosion.
Step 2.5: Replace main housing oil seal rings with Part # T202-SEAL-KIT.

4. POST-MAINTENANCE SIGN-OFF & OPERATIONAL RE-START
--------------------------------------------------------------------------------
- Inspector Name: Lead Maintenance Engineer (admin@samiq.ai)
- Completion Criteria: Zero active thermal alarms, vibration < 1.8 mm/s at 3,600 RPM.
- Authorized Sign-off: Chief Operations Manager
================================================================================
"""

with open(txt_path, "w", encoding="utf-8") as f:
    f.write(txt_content)
print(f"Created: {txt_path}")

# ---------------------------------------------------------
# 2. Multi-page PDF Document (2 Pages)
# ---------------------------------------------------------
pdf_path = os.path.join(OUTPUT_DIR, "SOP_702_Centrifugal_Pump_Maintenance_Guide.pdf")
doc = SimpleDocTemplate(pdf_path, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
styles = getSampleStyleSheet()

title_style = ParagraphStyle(
    'DocTitle',
    parent=styles['Heading1'],
    fontName='Helvetica-Bold',
    fontSize=18,
    leading=22,
    textColor=colors.HexColor('#0F4C81'),
    spaceAfter=12
)

subtitle_style = ParagraphStyle(
    'DocSubTitle',
    parent=styles['Heading2'],
    fontName='Helvetica-Bold',
    fontSize=13,
    leading=16,
    textColor=colors.HexColor('#10B981'),
    spaceAfter=10
)

body_style = ParagraphStyle(
    'DocBody',
    parent=styles['Normal'],
    fontName='Helvetica',
    fontSize=10,
    leading=14,
    textColor=colors.HexColor('#1E293B'),
    spaceAfter=8
)

story = []

# PAGE 1
story.append(Paragraph("SOP-702: Centrifugal Pump (P-101) Maintenance Standard", title_style))
story.append(Paragraph("Document Ref: SOP-702-REV4 | Target Tag: P-101 | Department: Operations", subtitle_style))
story.append(Spacer(1, 10))

story.append(Paragraph("<b>1. SECTION OVERVIEW & EQUIPMENT SPECIFICATIONS</b>", subtitle_style))
story.append(Paragraph("This manual specifies operational safety and preventive maintenance routines for Centrifugal Water Pump <b>P-101</b> installed in primary cooling loop A.", body_style))

# Table Specs
data_p1 = [
    ['Specification Parameter', 'Operating Bound', 'Alert Threshold'],
    ['Design Flow Rate', '450 GPM', 'Min 380 GPM'],
    ['Bearing Temperature Limit', '75.0 °C', 'Max Alarm 82.0 °C'],
    ['Discharge Pressure', '120 PSIG', 'Min 95 PSIG'],
    ['Recommended Lubricant', 'Mobil SHC 630', 'Change Every 2,000 Hrs']
]
t1 = Table(data_p1, colWidths=[200, 150, 150])
t1.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0F4C81')),
    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ('FONTSIZE', (0,0), (-1,0), 9),
    ('BOTTOMPADDING', (0,0), (-1,0), 6),
    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
    ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#F8FAFC')),
]))
story.append(t1)
story.append(Spacer(1, 15))

story.append(Paragraph("<b>2. BEARING LUBRICATION & SEAL INSPECTION PROCEDURE</b>", subtitle_style))
story.append(Paragraph("1. Verify oil sump sight glass level is at exactly 50% midpoint.", body_style))
story.append(Paragraph("2. Flush contaminated grease via mechanical drain plug until clean lubricant emerges.", body_style))
story.append(Paragraph("3. Inspect mechanical face seal for leakage rates exceeding 5 drops per minute.", body_style))

# End of Page 1 -> Force PageBreak
story.append(PageBreak())

# PAGE 2
story.append(Paragraph("<b>3. TROUBLESHOOTING & FAILURE DIAGNOSTICS (PAGE 2)</b>", title_style))
story.append(Paragraph("Refer to the decision matrix below when pump P-101 triggers thermal or vibration telemetry alarms.", body_style))
story.append(Spacer(1, 10))

data_p2 = [
    ['Observed Fault', 'Probable Cause', 'Corrective SOP Action'],
    ['Bearing Temp > 80°C', 'Oil Degradation / Contamination', 'Execute Lubricant Flush SOP-441'],
    ['Vibration > 3.0 mm/s', 'Coupling Misalignment / Wear', 'Re-align Motor Shaft with Laser Tool'],
    ['Discharge Flow Drop', 'Impeller Clogging / Cavitation', 'Inspect Intake Suction Strainer']
]
t2 = Table(data_p2, colWidths=[150, 175, 175])
t2.setStyle(TableStyle([
    ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#10B981')),
    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ('FONTSIZE', (0,0), (-1,0), 9),
    ('BOTTOMPADDING', (0,0), (-1,0), 6),
    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
    ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#F1F5F9')),
]))
story.append(t2)
story.append(Spacer(1, 15))

story.append(Paragraph("<b>4. PREVENTIVE MAINTENANCE SIGN-OFF LOG</b>", subtitle_style))
story.append(Paragraph("Authorized Technician: <b>employee@samiq.ai</b><br/>Reviewed by Operations Lead: <b>admin@samiq.ai</b><br/>Date Executed: <b>2026-07-21</b>", body_style))

doc.build(story)
print(f"Created: {pdf_path}")

# ---------------------------------------------------------
# 3. Multi-page DOCX Document (2 Pages)
# ---------------------------------------------------------
docx_path = os.path.join(OUTPUT_DIR, "Incident_Report_Reciprocating_Compressor_C303.docx")
d = docx.Document()

# PAGE 1
d.add_heading("INCIDENT & INSPECTION REPORT: COMPRESSOR C-303", level=1)
d.add_paragraph("Document ID: INC-2026-C303 | Equipment Tag: C-303 | Department: Maintenance")
d.add_paragraph("----------------------------------------------------------------------------------")

d.add_heading("1. Incident Summary & Alarm Logs", level=2)
d.add_paragraph(
    "On July 21, 2026, Reciprocating Air Compressor C-303 triggered an automated inspection warning "
    "due to elapsed operational hours exceeding the scheduled 500-hour oil check interval by 12 hours."
)

table1 = d.add_table(rows=1, cols=3)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Log Timestamp'
hdr_cells[1].text = 'Severity'
hdr_cells[2].text = 'Trigger Description'

row_cells = table1.add_row().cells
row_cells[0].text = '2026-07-21 10:05 AM'
row_cells[1].text = 'Info / Warning'
row_cells[2].text = 'Lubrication Inspection Overdue (C-303)'

d.add_paragraph().paragraph_format.space_after = docx.shared.Pt(14)
d.add_paragraph(
    "Preliminary diagnostic evaluation indicates zero thermal runaway. Compressor valve pressures "
    "remain within normal baseline (140 PSIG)."
)

# Page break for Page 2
d.add_page_break()

# PAGE 2
d.add_heading("2. Corrective Action Plan & Spare Parts (Page 2)", level=1)
d.add_paragraph(
    "Maintenance technician admin@samiq.ai conducted visual and fluid checks on Compressor C-303."
)

d.add_heading("Recommended Maintenance Actions:", level=2)
d.add_paragraph("• Step 1: Drain 15 Liters of synthetic lubricant from primary crankcase sump.")
d.add_paragraph("• Step 2: Replace oil filter element Part # C303-FILT-09.")
d.add_paragraph("• Step 3: Torque cylinder head bolts to 180 ft-lbs.")
d.add_paragraph("• Step 4: Re-certify baseline pressure differential and update telemetry log.")

d.add_heading("Safety & Sign-off Certification:", level=2)
d.add_paragraph("Lead Technician: admin@samiq.ai\nDepartment Head: Maintenance Supervisor\nStatus: Action Pending Completion")

d.save(docx_path)
print(f"Created: {docx_path}")

# ---------------------------------------------------------
# 4. Multi-page Technical Image Document (.PNG)
# ---------------------------------------------------------
img_path = os.path.join(OUTPUT_DIR, "P101_Pump_Electrical_Schematic.png")
width, height = 1000, 700
image = Image.new('RGB', (width, height), color='#0F172A')
draw = ImageDraw.Draw(image)

# Draw Title & Border
draw.rectangle([(20, 20), (width-20, height-20)], outline='#10B981', width=3)
draw.rectangle([(30, 30), (width-30, 90)], fill='#1E293B', outline='#334155', width=1)
draw.text((40, 45), "EQUIPMENT DIAGRAM: P-101 MOTOR CONTROL & TELEMETRY CIRCUIT", fill='#FFFFFF')
draw.text((40, 65), "Tag: P-101 | Department: Operations | Drawing Ref: ELEC-P101-SHEET-1", fill='#10B981')

# Draw Equipment Blocks & Circuits
draw.rectangle([(100, 160), (320, 300)], fill='#1E293B', outline='#10B981', width=2)
draw.text((120, 180), "[ POWER SUPPLY ]\n480V 3-Phase MCC\nBreaker: CB-101", fill='#FFFFFF')

draw.line([(320, 230), (450, 230)], fill='#3B82F6', width=4)
draw.text((350, 210), "Feed Line", fill='#94A3B8')

draw.rectangle([(450, 160), (700, 300)], fill='#1E293B', outline='#3B82F6', width=2)
draw.text((470, 180), "[ MAIN MOTOR (P-101) ]\n75 HP Centrifugal Pump\nSpeed: 1,750 RPM", fill='#FFFFFF')

# Draw Telemetry Sensors
draw.line([(570, 300), (570, 450)], fill='#F59E0B', width=3)

draw.rectangle([(450, 450), (700, 560)], fill='#1E293B', outline='#F59E0B', width=2)
draw.text((470, 470), "[ SENSOR SUITE ]\nTemp Probe PT-101: 72°C\nVib Sensor VT-101: 1.4 mm/s", fill='#F59E0B')

# Footer Info
draw.text((40, 630), "SamiQ Industrial AI Knowledge Ingestion File | Verified Technical Drawing", fill='#64748B')

image.save(img_path)
print(f"Created: {img_path}")

print("All 4 multi-page industrial test documents generated successfully!")
